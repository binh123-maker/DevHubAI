"""DOCX text extraction using python-docx."""
from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from app.services.processors.base import ExtractedDocument, ExtractedPage


def _paragraph_to_markdown(para) -> str:  # type: ignore[type-arg]
    """
    Convert a single docx paragraph to a Markdown snippet.

    Supported styles: Heading 1-6 → # to ######, List Paragraph → bullet.
    Everything else is rendered as plain text.
    """
    style_name: str = para.style.name if para.style else ""
    text: str = para.text.strip()

    if not text:
        return ""

    # Headings
    if style_name.startswith("Heading 1"):
        return f"# {text}"
    if style_name.startswith("Heading 2"):
        return f"## {text}"
    if style_name.startswith("Heading 3"):
        return f"### {text}"
    if style_name.startswith("Heading 4"):
        return f"#### {text}"
    if style_name.startswith("Heading 5"):
        return f"##### {text}"
    if style_name.startswith("Heading 6"):
        return f"###### {text}"

    # List items
    if style_name.startswith("List"):
        return f"- {text}"

    # Bold / italic via runs
    md_runs: list[str] = []
    for run in para.runs:
        run_text = run.text
        if not run_text:
            continue
        if run.bold and run.italic:
            run_text = f"***{run_text}***"
        elif run.bold:
            run_text = f"**{run_text}**"
        elif run.italic:
            run_text = f"*{run_text}*"
        md_runs.append(run_text)

    return "".join(md_runs) if md_runs else text


def extract_docx(file_path: Path) -> ExtractedDocument:
    """
    Extract text from a DOCX file.

    DOCX has no native page concept — we group every 30 paragraphs into a
    synthetic 'page' so that chunk metadata remains meaningful.
    """
    docx = DocxDocument(str(file_path))
    paragraphs = [p for p in docx.paragraphs if p.text.strip()]

    PAGE_SIZE = 30  # paragraphs per synthetic page
    pages: list[ExtractedPage] = []

    for page_idx, start in enumerate(range(0, len(paragraphs), PAGE_SIZE), start=1):
        chunk_paras = paragraphs[start : start + PAGE_SIZE]

        raw_lines = [p.text.strip() for p in chunk_paras if p.text.strip()]
        md_lines = [_paragraph_to_markdown(p) for p in chunk_paras if p.text.strip()]

        raw_text = "\n".join(raw_lines)
        markdown = "\n\n".join(md_lines)

        if raw_text:
            pages.append(ExtractedPage(
                page_number=page_idx,
                raw_text=raw_text,
                markdown=markdown,
            ))

    return ExtractedDocument(pages=pages)
