import uuid
import time
import os
import hashlib
from pathlib import Path
from sqlalchemy import select
from sqlalchemy.orm import Session
import fitz
import docx

from app.models.document import Document, DocumentBinary, DocumentVersion, DocumentStructureNode, DocumentChunk, ProcessingJob
from app.models.user import User, UserRole
from app.models.workspace import Workspace
from app.services.processing_service import execute_processing_job

def test_integration_pipeline_formats_and_metadata(db_session: Session):
    # 1. Setup User and Workspace
    user = User(
        email=f"integ_{uuid.uuid4().hex[:8]}@example.com",
        password_hash="password",
        role=UserRole.USER,
        is_active=True
    )
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)

    workspace = Workspace(name="Integ Test WS", user_id=user.id)
    db_session.add(workspace)
    db_session.commit()
    db_session.refresh(workspace)

    # Helper function to run pipeline and assert
    def run_pipeline_for_file(file_path: Path, file_type: str):
        doc = Document(
            user_id=user.id,
            workspace_id=workspace.id,
            title=f"Test {file_type.upper()}",
            file_name=file_path.name,
            file_type=file_type,
            file_size=file_path.stat().st_size,
            file_path=str(file_path.absolute()),
            status="uploading"
        )
        db_session.add(doc)
        db_session.commit()
        db_session.refresh(doc)

        binary = DocumentBinary(
            sha256=f"hash_{uuid.uuid4().hex}",
            file_path=str(file_path.absolute()),
            file_type=file_type,
            file_size=file_path.stat().st_size
        )
        db_session.add(binary)
        db_session.commit()
        db_session.refresh(binary)

        version = DocumentVersion(
            document_id=doc.id,
            version_number=1,
            binary_id=binary.sha256,
            status="uploading"
        )
        db_session.add(version)
        db_session.commit()
        db_session.refresh(version)

        job = ProcessingJob(
            document_version_id=version.id,
            job_type="upload_processing",
            status="pending",
            progress=0,
            retry_count=0
        )
        db_session.add(job)
        db_session.commit()
        db_session.refresh(job)

        # Trigger execution
        execute_processing_job(db_session, job.id)
        
        # Reload status
        db_session.refresh(job)
        db_session.refresh(version)
        db_session.refresh(doc)

        return doc, version, job

    # 2. PDF Processing Test
    pdf_path = Path("test_doc.pdf")
    doc_pdf = fitz.open()
    page = doc_pdf.new_page()
    page.insert_text((50, 50), "# PDF Heading 1\nThis is text from PDF paragraph.")
    doc_pdf.save(str(pdf_path.absolute()))
    doc_pdf.close()

    try:
        doc, version, job = run_pipeline_for_file(pdf_path, "pdf")
        assert job.status == "completed"
        
        # Verify chunks saved
        chunks = db_session.scalars(select(DocumentChunk).where(DocumentChunk.document_version_id == version.id).order_by(DocumentChunk.chunk_index)).all()
        assert len(chunks) > 0
        assert chunks[0].chunk_index == 0
        assert "estimated_tokens" in chunks[0].metadata_json
        
    finally:
        if pdf_path.exists():
            pdf_path.unlink()

    # 3. DOCX Processing Test
    docx_path = Path("test_doc.docx")
    doc_docx = docx.Document()
    doc_docx.add_heading("Docx Heading 1", level=1)
    doc_docx.add_paragraph("This is docx paragraph content text.")
    doc_docx.save(str(docx_path.absolute()))

    try:
        doc, version, job = run_pipeline_for_file(docx_path, "docx")
        assert job.status == "completed"
        
        chunks = db_session.scalars(select(DocumentChunk).where(DocumentChunk.document_version_id == version.id)).all()
        assert len(chunks) > 0
        
    finally:
        if docx_path.exists():
            docx_path.unlink()

    # 4. Markdown Processing Test
    md_path = Path("test_doc.md")
    md_path.write_text("# Markdown Title\n\nSome body text here.\n\n## Sub heading\nMore details text.", encoding="utf-8")

    try:
        doc, version, job = run_pipeline_for_file(md_path, "md")
        assert job.status == "completed"
        
        chunks = db_session.scalars(select(DocumentChunk).where(DocumentChunk.document_version_id == version.id)).all()
        assert len(chunks) > 0
        
        # Verify metadata propagation
        meta = chunks[0].metadata_json
        assert "heading_path" in meta
        assert "importance_score" in meta
        assert "retrieval_weight" in meta
        
    finally:
        if md_path.exists():
            md_path.unlink()

    # 5. TXT Processing Test
    txt_path = Path("test_doc.txt")
    txt_path.write_text("Hello from plaintext file. This should be chunked.", encoding="utf-8")

    try:
        doc, version, job = run_pipeline_for_file(txt_path, "txt")
        assert job.status == "completed"
        
        chunks = db_session.scalars(select(DocumentChunk).where(DocumentChunk.document_version_id == version.id)).all()
        assert len(chunks) > 0
        
    finally:
        if txt_path.exists():
            txt_path.unlink()


def test_empty_document_handling(db_session: Session):
    # Setup Workspace, Document, version, job
    user = User(
        email=f"empty_{uuid.uuid4().hex[:8]}@example.com",
        password_hash="password",
        role=UserRole.USER,
        is_active=True
    )
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)

    workspace = Workspace(name="Empty Test WS", user_id=user.id)
    db_session.add(workspace)
    db_session.commit()
    db_session.refresh(workspace)

    empty_path = Path("empty.txt")
    empty_path.write_text("", encoding="utf-8")

    try:
        doc = Document(
            user_id=user.id,
            workspace_id=workspace.id,
            title="Empty Doc",
            file_name=empty_path.name,
            file_type="txt",
            file_size=0,
            file_path=str(empty_path.absolute()),
            status="uploading"
        )
        db_session.add(doc)
        db_session.commit()
        db_session.refresh(doc)

        binary = DocumentBinary(
            sha256=f"hash_{uuid.uuid4().hex}",
            file_path=str(empty_path.absolute()),
            file_type="txt",
            file_size=0
        )
        db_session.add(binary)
        db_session.commit()
        db_session.refresh(binary)

        version = DocumentVersion(
            document_id=doc.id,
            version_number=1,
            binary_id=binary.sha256,
            status="uploading"
        )
        db_session.add(version)
        db_session.commit()
        db_session.refresh(version)

        job = ProcessingJob(
            document_version_id=version.id,
            job_type="upload_processing",
            status="pending",
            progress=0,
            retry_count=0
        )
        db_session.add(job)
        db_session.commit()
        db_session.refresh(job)

        execute_processing_job(db_session, job.id)
        db_session.refresh(job)
        
        # Empty document should succeed or handle gracefully (completed)
        assert job.status in ("completed", "failed")
        
    finally:
        if empty_path.exists():
            empty_path.unlink()
